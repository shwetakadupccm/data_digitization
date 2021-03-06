import shutil
import math
import os
import re
import pandas as pd
import numpy as np
import pyqrcode
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pytesseract as pt
from pdf2image import convert_from_path
from PyPDF2 import PdfFileReader, PdfFileWriter
from docx2pdf import convert

pt.pytesseract.tesseract_cmd = 'C:/Program Files/Tesseract-OCR/tesseract.exe'

class DataDigitization():

    def __init__(self, root, tmp_folder_path, destination_path):
        self.root = root
        self.tmp_folder_path = tmp_folder_path
        self.report_names_df = pd.read_excel(os.path.join(self.root, 'reference_docs/Report_types_17.xlsx'))
        self.categorized_files_df = pd.read_excel(os.path.join(self.root, 'reference_docs/', '2022_02_09_Data_digitzation_scanned_files_dk.xlsx'))
        self.tmp_folder_path = tmp_folder_path
        self.scanned_patient_file_path = os.path.join(self.root, 'scanned_patient_files/2022_03_14/original_pdf')
        self.scanned_files_split_path = os.path.join(self.root, 'scanned_patient_files/2022_03_15')
        self.categorized_files_path = os.path.join(self.root, 'sample_output/2022_03_14/categorized_files_path')
        self.destination_path = destination_path

    @staticmethod
    def make_qr_code(file_number, mr_number, report_type, subfolder, destination):
        file_number_str = re.sub('_', '/',
                                 str(file_number))
        if subfolder is not None:
            qr_code = file_number_str + '_' + \
                str(mr_number) + '_' + str(report_type) + '_' + str(subfolder)
        else:
            qr_code = file_number_str + '_' + \
                str(mr_number) + '_' + str(report_type)
        qr = pyqrcode.create(qr_code)
        report_type_for_name = re.sub(' ', '_', str(report_type))
        qr_img_name = file_number + '_' + \
            str(mr_number) + '_' + report_type_for_name + '.png'
        qr_path = os.path.join(destination, qr_img_name)
        qr.png(qr_path, scale=4)
        print('QR code created for ' + file_number + ' ' + report_type + ' ')
        return qr_img_name

    def add_qr_code_in_word_doc(self, report_type, qr_code_path, file_number, mr_number,
                                patient_name, dob):
        doc = Document()
        doc.add_picture(qr_code_path)
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        type = report_type
        text = doc.add_paragraph()
        report_type_name = text.add_run(str(type))
        report_type_name.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # type: ignore
        report_type_name.bold = True
        report_type_name.font.size = Pt(28)
        report_type_name.font.name = 'Arial Black'
        blank_para = doc.add_paragraph()
        run = blank_para.add_run()
        run.add_break()
        file_number_str = re.sub('_', '/', str(file_number))
        file_no = 'File Number: ' + str(file_number_str)
        id = doc.add_paragraph().add_run(file_no)
        id.alignment = WD_ALIGN_PARAGRAPH.CENTER
        id.font.size = Pt(20)
        id.font.name = 'Arial Black'
        mr_no = 'MR Number: ' + str(mr_number)
        id = doc.add_paragraph().add_run(mr_no)
        id.font.size = Pt(20)
        id.font.name = 'Arial Black'
        pt_name = 'Patient Name: ' + patient_name
        id = doc.add_paragraph().add_run(pt_name)
        id.font.size = Pt(20)
        id.font.name = 'Arial Black'
        date_of_birth = 'Date of Birth: ' + dob
        id = doc.add_paragraph().add_run(date_of_birth)
        id.font.size = Pt(20)
        id.font.name = 'Arial Black'
        report_type = re.sub(' ', '_', report_type)
        report_type = report_type.lower()
        doc_name = str(file_number) + '_' + str(mr_number) + \
            '_' + str(report_type) + '.docx'
        coded_data = os.path.join(self.tmp_folder_path, 'coded_data')
        if not os.path.isdir(coded_data):
            os.mkdir(coded_data)
        doc_path = os.path.join(coded_data, doc_name)
        doc.save(doc_path)
        return doc_path

    @staticmethod
    def convert_doc_to_pdf(doc_path):
        pdf_path = re.sub('.docx', '.pdf', str(doc_path))
        convert(doc_path, pdf_path)
        return pdf_path

    @staticmethod
    def split_pdf_by_pages(file_number, scanned_files_path, splitted_file_path):
        pdf_file_name = str(file_number) + '.pdf'
        scanned_file = PdfFileReader(os.path.join(scanned_files_path, pdf_file_name))
        page_range = scanned_file.getNumPages()
        for i in range(page_range):
            page = scanned_file.getPage(i)
            page_no = i + 1
            splitted_file = str(file_number) + '_' + str(page_no) + '.pdf'
            pdf_writer = PdfFileWriter()
            pdf_writer.addPage(page)
            with open(os.path.join(splitted_file_path, splitted_file), 'wb') as out:
                pdf_writer.write(out)
        print("file number: ", file_number + " splitted")

    @staticmethod
    def get_image_no(file_number, file_images_lst):
        file_images_no_lst = []
        for file_image in file_images_lst:
            file_image_no = re.sub(file_number, '', str(file_image))
            file_image_no = re.sub('.pdf', '', file_image_no)
            file_image_no = re.sub('_', '', file_image_no)
            file_image_no = file_image_no.strip()
            file_images_no_lst.append(file_image_no)
        return file_images_no_lst

    @staticmethod
    def split_report_page_no(report_page_no):
        if isinstance(report_page_no, float):
            page_no_lst = []
            if not math.isnan(report_page_no):
                integer = int(report_page_no)
                page_no_lst.append(str(integer))
            return page_no_lst
        elif isinstance(report_page_no, int):
            page_no_lst = []
            page_no_lst.append(str(report_page_no))
            return page_no_lst
        elif ';' in report_page_no:
            page_no_lst = []
            report_page_no_splitted = report_page_no.split(';')
            for page_no in report_page_no_splitted:
                if '|' in page_no:
                    partitions = page_no.partition('|')
                    start = int(partitions[0])
                    end = int(partitions[2]) + 1
                    page_nos = np.arange(start, end)
                    page_nos_lst = page_nos.tolist()
                    for no in page_nos_lst:
                        page_no_lst.append(str(no))
                else:
                    page_no_lst.append(str(page_no))
            return page_no_lst
        elif '|' in report_page_no:
            page_no_lst = []
            partitions = report_page_no.partition('|')
            start = int(partitions[0])
            end = int(partitions[2]) + 1
            page_nos = np.arange(start, end)
            page_nos_lst = page_nos.tolist()
            for no in page_nos_lst:
                page_no_lst.append(str(no))
            return page_no_lst
        elif type(report_page_no) in (float, int):
            page_no_lst = []
            report_page_no = int(report_page_no)
            page_no_lst.append(str(report_page_no))
            return page_no_lst
        else:
            page_no_lst = []
            page_no_lst.append(str(report_page_no))
            return page_no_lst

    @classmethod
    def classify_file_images_by_report_types(self, splitted_file_path_file_no, report_page_nums, file_number,
                                             report_type, destination_path):
        splitted_scanned_files = os.listdir(splitted_file_path_file_no)
        img_no_lst = self.get_image_no(file_number, splitted_scanned_files)
        report_page_no_splitted = self.split_report_page_no(report_page_nums)
        file_no_dir = os.path.join(destination_path, str(file_number))
        if not os.path.isdir(file_no_dir):
            os.mkdir(file_no_dir)
        report_dir = os.path.join(file_no_dir, report_type)
        if not os.path.isdir(report_dir):
            os.mkdir(report_dir)
        for page_no in report_page_no_splitted:
            if page_no in img_no_lst:
                report_page_name = str(file_number) + \
                    '_' + str(page_no) + '.pdf'
                source_path = os.path.join(
                    splitted_file_path_file_no, report_page_name)
                dest_path = os.path.join(report_dir, report_page_name)
                shutil.copy(source_path, dest_path)

    @staticmethod
    def rename_images(pdf_doc_path, dir_path, file_no, report_type, destination_path):
        report_dir = os.path.join(dir_path, str(report_type))
        img_list = os.listdir(report_dir)
        for index, img in enumerate(img_list):
            old_file_path = os.path.join(report_dir, img)
            img_no = index + 1
            new_name = str(file_no) + '_' + str(img_no) + '.pdf'
            file_dir = os.path.join(destination_path, str(file_no))
            if not os.path.isdir(file_dir):
                os.mkdir(file_dir)
            new_file_path = os.path.join(file_dir, report_type)
            if not os.path.isdir(new_file_path):
                os.mkdir(new_file_path)
            dest_path = os.path.join(new_file_path, new_name)
            shutil.copy(old_file_path, dest_path)
            coded_file_name = 'code_' + \
                str(file_no) + '_' + str(report_type) + '.pdf'
            shutil.copy(pdf_doc_path, os.path.join(
                new_file_path, coded_file_name))

    def categorize_file_by_report_types(self):
        for i in range(len(self.categorized_files_df)):
            file_number = self.categorized_files_df['file_number'][i]
            mr_number = self.categorized_files_df['mr_number'][i]
            patient_name = self.categorized_files_df['patient_name'][i]
            dob = self.categorized_files_df['date_of_birth'][i]
            splitted_files_dir = os.path.join(self.tmp_folder_path, 'splitted_files')
            if not os.path.isdir(splitted_files_dir):
                os.mkdir(splitted_files_dir)
            splitted_file_dir = os.path.join(splitted_files_dir, str(file_number))
            if not os.path.isdir(splitted_file_dir):
                os.mkdir(splitted_file_dir)
            self.split_pdf_by_pages(str(file_number), self.scanned_patient_file_path, splitted_file_dir)
            for report_type in self.report_names_df['report_number_and_type']:
                qr_code_dir = os.path.join(self.tmp_folder_path, 'qr_codes')
                if not os.path.isdir(qr_code_dir):
                    os.mkdir(qr_code_dir)
                qr_img_name = self.make_qr_code(
                    file_number, mr_number, report_type, None, qr_code_dir)
                coded_data_dir = os.path.join(
                    self.tmp_folder_path, 'coded_data')
                if not os.path.isdir(coded_data_dir):
                    os.mkdir(coded_data_dir)
                qr_code_path = os.path.join(qr_code_dir, qr_img_name)
                coded_doc_path = self.add_qr_code_in_word_doc(
                    report_type, qr_code_path, file_number, mr_number,
                    patient_name, dob)
                coded_pdf_path = self.convert_doc_to_pdf(coded_doc_path)
                report_type_str = re.sub(' ', '_', str(report_type))
                report_page_nums = self.categorized_files_df[report_type_str][i]
                classified_files_path = os.path.join(
                    self.tmp_folder_path, 'classfied_files')
                if not os.path.isdir(classified_files_path):
                    os.mkdir(classified_files_path)
                self.classify_file_images_by_report_types(splitted_file_dir, str(
                    report_page_nums), file_number, report_type_str, classified_files_path)
                renamed_files_path = os.path.join(
                    classified_files_path, str(file_number))
                self.rename_images(coded_pdf_path, renamed_files_path, str(
                    file_number), report_type_str, self.destination_path)
                print("file: ", file_number,
                      ' classified by report types and arranged by sequence')
