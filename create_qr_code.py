import pandas as pd
import os
import re
import pyqrcode
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pytesseract as pt
from PyPDF2 import PdfFileReader, PdfFileWriter
# from PIL import Image
from pdf2image import convert_from_path
import shutil
from docx2pdf import convert
from file_functions import HelperFunctions


class QrCode():

    def __init__(self, root, master_list_name, categorized_excel):
        self.root = root
        self.master_list_name = master_list_name
        self.categorized_excel = categorized_excel
        self.hf = HelperFunctions(
            self.root, self.master_list_name, self.categorized_excel)

    def add_qr_code_in_word_document(self, category_row):
        qr_img_path, qr_code_lst = self.make_qr_code(category_row)
        print(qr_code_lst)
        doc = Document()
        doc.add_picture(qr_img_path)
        qr_alignment = doc.paragraphs[-1]
        qr_alignment.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        folders = qr_code_lst[4:]
        folder = [str(folder) for folder in folders if folder != 'nan']
        report_type = ' '.join(str(id) for id in folder)
        doc = self.hf.format_word_doc(report_type, doc)
        blank_para = doc.add_paragraph()
        run = blank_para.add_run()
        run.add_break()
        prefixes = self.hf.function_params('qr_data_cols')
        for idx, prefix in enumerate(prefixes):
            doc = self.hf.format_word_doc(
                str(prefix) + ': ' + str(qr_code_lst[idx]), doc)
        doc_name = 'coded_file.docx'
        coded_data_dir = self.create_tmp_folder_for_data_type(
            data_type='coded_data')
        doc_path = os.path.join(os.path.join(coded_data_dir, doc_name))
        doc.save(doc_path)
        pdf_path = doc_path.replace('.docx', '.pdf')
        convert(doc_path, pdf_path)
        return qr_code_lst, pdf_path

    def create_tmp_folder_for_data_type(self, data_type):
        data_type_dir = os.path.join(self.root, 'tmp/', data_type)
        if not os.path.isdir(data_type_dir):
            os.mkdir(data_type_dir)
        return data_type_dir

    @classmethod
    def make_patient_full_name(self, master_list):
        first_name_col_idx = master_list.columns.get_loc('first_name')
        last_name_col_idx = master_list.columns.get_loc('last_name')
        master_list['patient_name'] = master_list[master_list.columns[first_name_col_idx:last_name_col_idx]].apply(
            lambda x: ' '.join(x.dropna().astype(str)), axis=1)
        master_list['patient_name'] = master_list['patient_name'].str.title()
        return master_list

    def get_id_data(self, file_number):
        """
        get id values from input id names and single row of master list
        :param master_list: pd.DataFrame
        :param file_number: file_number separated by '_'
        :param id_cols: col-names which stores the patients identifing info(file_number, mr_number, name, dob)
        :return: list of id data for single row of master list
        """
        master_list = self.hf.function_params('master_list')
        master_list = self.make_patient_full_name(master_list)
        id_cols = self.hf.function_params('id_cols')
        id_data = master_list[master_list['file_number']
                              == file_number][id_cols]

        id_data_lst = id_data.values.tolist()
        id_data = [id_text for sublist in id_data_lst for id_text in sublist]
        return id_data

    def make_qr_code(self, category_row):
        """
        create qr code from input data with proper formating. returns the data used
        to create the qr code
        :param master_list_row:
        :param category_row:
        :param qr_destination_path:
        :return: qr_code_dat
        """
        folders = list(
            category_row[self.hf.function_params('folder_col_heads')])
        file_number = category_row[self.hf.function_params(
            'file_number')].get('file_number')
        id_dat = self.get_id_data(file_number)
        file_number_str = file_number.replace("_", "/")
        folder = [str(folder) for folder in folders if folder != 'nan']
        id_data = [file_number_str, str(id_dat[0])] + folder
        qr_code_dat = '_'.join(str(id_text) for id_text in id_data)
        qr_code_dat = re.sub('_nan', '', str(qr_code_dat))
        qr_code_lst = [file_number_str] + id_dat + folder
        qr_img = pyqrcode.create(qr_code_dat)
        qr_dir = self.create_tmp_folder_for_data_type(data_type='qr_code')
        qr_img_path = os.path.join(qr_dir, 'qr_img.png')
        qr_img.png((qr_img_path), scale=4)
        return qr_img_path, qr_code_lst
