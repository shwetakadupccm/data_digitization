'''
functions that are required to move/rename files create folders etc
define paramters
'''
import os
import re
import math
import pandas as pd
import numpy as np
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH
from docx.shared import Pt

REPORT_TYPE_DICT = dict(patient_information=1,
                        clinical_examination=2,
                        radiology=3,
                        metastatic_examination=4,
                        biopsy_pathology=5,
                        neo_adjuvant_chemotherapy=6,
                        surgical_rocedures=7,
                        patient_images=8,
                        surgery_media=9,
                        surgery_pathology=10,
                        chemotherapy=11,
                        radiotherapy=12,
                        follow_up_notes=13,
                        genetics=14,
                        miscellaneous=15,
                        patient_file_data=16,
                        PROMS=17)


class HelperFunctions:
    '''
    functions that are required to move/rename files create folders etc
    define paramters
    '''

    # def __init__(self):
    # self.root = root
    # self.master_list_name = master_list_name
    # self.categorized_excel = categorized_excel
    # ignore type unexpected-keyword-arg)

    def function_params(self, param_name):
        '''
        define paths and names and data dfs used for all data_digitization functions
        '''
        param_dict = dict(status_sheet=pd.read_excel(
            os.path.join('status_sheets', '2019_2022_scanned_files_status_sheet.xlsx')),
            id_cols=['mr_number',
                     'patient_name', 'date_of_birth'],
            file_number=['file_number'],
            qr_data_cols=['File Number', 'MR Number',
                          'Patient Name', 'Date of Birth'],
            folder_col_heads=['report_name', 'subfolder_name'],
            report_types_list=REPORT_TYPE_DICT)
        dat = param_dict.get(param_name)
        return dat

    # def create_folder_for_data_type(self, data_type, source_folder):
    #     '''
    #     create folders and subfolders for data types in source folders for eg. tmp
    #     '''
    #     data_type_dir = False
    #     if not os.path.isdir(os.path.join(self.root, source_folder)):
    #         os.mkdir(os.path.join(self.root, source_folder))
    #     dat = []
    #     for folder in data_type:
    #         dat.append(folder)
    #         folders = "/".join(dat)
    #         data_type_dir = os.path.join(self.root, source_folder, folders)
    #         if not os.path.isdir(data_type_dir):
    #             os.mkdir(data_type_dir)
    #     return data_type_dir

    # data type = list?
    def create_folder_for_data_type(self, data_type, source_folder):
        source_path = os.path.join(self.root, source_folder)
        if not os.path.isdir(source_path):
            os.mkdir(source_path)
        data_type_dir = os.path.join(source_path, data_type)
        if not os.path.isdir(data_type_dir):
            os.mkdir(data_type_dir)
        return data_type_dir

    def create_categorized_df(self):
        '''
        creates list of file numbers that have not been categorized and fetch categorization data from 
        categorization excels
        '''
        # to do need to create df from multiple user rows..so pick up fron each user categorized_data data by file number for that user
        status_df = self.function_params('status_sheet')
        tbd_rows = status_df[status_df.categorized_status == 'yes']
        cat_by = tbd_rows.categorized_by.unique().tolist()
        tbd_df = tbd_rows['file_number', 'categorized_by'].drop_duplicates()
        category_df = pd.DataFrame()
        for file_name in os.listdir(os.path.join(self.root, 'categorization_data')):
            if re.match('file_categorization', file_name):
                categorized = pd.read_excel(
                    os.path.join('categorization_data', file_name))
                for cat in cat_by:
                file_tbd = tbd_df[tbd_df.categorized_by ==
                                  'cat_by'].unique().tolist()
                file_tbd = tbd_df[tbd_df.categorized_by == categorized_by][file_number].drop_duplicates(
                    inplace=True).tolist()
                cat_df = pd.merge(file_tbd, categorized,
                                  how='left', on='file_number', validate='1:m')
                category_df = pd.concat(category_df, cat_df, ignore_index=True)
        return category_df

    @staticmethod
    def change_sep(string, old_sep, new_sep):
        """
        change the separator between the string or within
        :param string: string
        :param old_sep: string separators (' ', '_', '/')
        :param new_sep: string separators (' ', '_', '/')
        :return: string with changed separators
        """
        changed_sep = re.sub(old_sep, new_sep, str(string))
        return changed_sep

    # @staticmethod
    # def split_report_page_no(report_page_no):
        # if isinstance(report_page_no, float):
        # page_no_lst = []
        # if not math.isnan(report_page_no):
        # integer = int(report_page_no)
        # page_no_lst.append(str(integer))
        # return page_no_lst
        # elif isinstance(report_page_no, int):
        # page_no_lst = []
        # page_no_lst.append(str(report_page_no))
        # return page_no_lst
        # elif ';' in report_page_no:
        # page_no_lst = []
        # report_page_no_splitted = report_page_no.split(';')
        # for page_no in report_page_no_splitted:
        # if '|' in page_no:
        # partitions = page_no.partition('|')
        # start = int(partitions[0])
        # end = int(partitions[2]) + 1
        # page_nos = np.arange(start, end)
        # page_nos_lst = page_nos.tolist()
        # for no in page_nos_lst:
        # page_no_lst.append(str(no))
        # else:
        # page_no_lst.append(str(page_no))
        # return page_no_lst
        # elif '|' in report_page_no:
        # page_no_lst = []
        # partitions = report_page_no.partition('|')
        # start = int(partitions[0])
        # end = int(partitions[2]) + 1
        # page_nos = np.arange(start, end)
        # page_nos_lst = page_nos.tolist()
        # for no in page_nos_lst:
        # page_no_lst.append(str(no))
        # return page_no_lst
        # elif type(report_page_no) in (float, int):
        # page_no_lst = []
        # report_page_no = int(report_page_no)
        # page_no_lst.append(str(report_page_no))
        # return page_no_lst
        # else:
        # page_no_lst = []
        # page_no_lst.append(str(report_page_no))
        # return page_no_lst

    @staticmethod
    def make_pdf_name_using_page_no(page_no, file_number):
        report_page_name = str(file_number) + '_' + str(page_no) + '.pdf'
        return report_page_name
    # @staticmethod
    # def format_word_doc(id_text, doc):
    #
    #     doc_text = doc.add_paragraph()
    #     doc_text = doc_text.add_run(id_text)
    #     doc_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #     doc_text.bold = True
    #     doc_text.font.size = Pt(12)
    #     doc_text.font.name = 'Arial Black'
    #     return doc

    @staticmethod
    def format_word_doc(id_text, doc):
        """
        format text strings in a word document
        :param id_text: id_value(eg = 'file_number', 'mr_number', 'patient_name', 'dob')
        :param doc: docx.Document
        :return: formatted document
        """
        doc_text = doc.add_paragraph()
        doc_text = doc_text.add_run(id_text)
        doc_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # type: ignore
        doc_text.bold = True
        doc_text.font.size = Pt(12)
        doc_text.font.name = 'Arial Black'
        return doc

# to do create check

    @staticmethod
    def split_report_page_no(report_page_numbers):
        page_list = None
        page_no_lst = str.split(report_page_numbers, ';')
        for page in page_no_lst:
            page_range = str.split(page, '|')
            if [page] != page_range:
                page_range = [int(page) for page in page_range]
                start_page = min(page_range)
                stop_page = max(page_range) + 1
                page_numbers = list(
                    range(start_page, stop_page))  # type: ignore
                if start_page != page_range[0]:
                    page_numbers.reverse()
                page = page_numbers
            else:
                page = int(page)
            if page_list is None:
                page_list = []
            if not isinstance(page, list):
                page = [page]
            page_list.append(page)
            print(page)
        print(page_list)
        page_list = [page for pages in page_list for page in pages]
        return page_list
